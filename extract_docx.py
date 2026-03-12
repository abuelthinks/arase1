import sys
import docx
import glob
import os

os.makedirs('extracted_forms', exist_ok=True)

docx_files = glob.glob('Updated Theruni Forms/*.docx')
for file_path in docx_files:
    if '~$' in file_path: continue
    try:
        doc = docx.Document(file_path)
        content = []
        for p in doc.paragraphs:
            if p.text.strip():
                content.append(p.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                content.append(' | '.join(row_data))
                
        out_name = os.path.join('extracted_forms', os.path.basename(file_path).replace('.docx', '.txt'))
        with open(out_name, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
        print(f'Extracted {file_path} to {out_name}')
    except Exception as e:
        print(f'Error extracting {file_path}: {e}')
